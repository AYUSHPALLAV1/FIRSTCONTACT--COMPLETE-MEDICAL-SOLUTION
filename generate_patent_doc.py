from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_patent_doc():
    doc = Document()

    # --- Helper Functions ---
    def add_section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_normal_paragraph(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Title ---
    title = doc.add_paragraph()
    run = title.add_run("PATENT SPECIFICATION DRAFT: FIRSTCONTACT")
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Digital Healthcare Accessibility and AI-Driven Preliminary Diagnosis")
    run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() # Spacer

    # --- 1. Field of Invention ---
    add_section_header("1. FIELD OF THE INVENTION")
    add_normal_paragraph("The present invention relates generally to the field of healthcare informatics and medical diagnostic systems. More specifically, it relates to an integrated digital platform that utilizes ensemble machine learning algorithms for preliminary disease prediction, automated triage, and holistic management of patient-doctor consultations and insurance policy analysis.")

    # --- 2. Background ---
    add_section_header("2. BACKGROUND OF THE INVENTION")
    add_normal_paragraph("Conventional healthcare systems often face significant delays in preliminary diagnosis and specialist referrals. Existing symptom checkers frequently operate in silos, providing generic medical information without considering a patient's vitals, such as Body Mass Index (BMI), Blood Pressure, and blood panel data. Furthermore, once a preliminary diagnosis is made, there is a disconnect between the diagnostic phase, the logistical phase (finding and booking a doctor), and the financial phase (determining insurance coverage). There is a critical need for an integrated system that automates the transition from symptom detection to actionable care.")

    # --- 3. Summary ---
    add_section_header("3. SUMMARY OF THE INVENTION")
    add_normal_paragraph("The present invention, titled \"FirstContact,\" is a comprehensive healthcare system that automates the triage process through an AI-driven engine. The system receives a plurality of patient inputs, including demographics, vital signs, and symptoms, and processes them through a Random Forest Classifier to generate a probabilistic diagnosis. The invention further includes a dynamic scheduling module for telemedicine, a geospatial hospital locator, and an automated text-analysis module for insurance policy evaluation, providing a unified solution for patient health management.")

    # --- 4. Technical Novelty ---
    add_section_header("4. TECHNICAL NOVELTY AND ADVANTAGES")
    novelties = [
        "Integrated Lifecycle Management: The primary novelty lies in the seamless integration of AI diagnosis with doctor booking and insurance analysis in a single session.",
        "Enhanced Accuracy via Ensemble Learning: Utilization of Random Forest algorithms reduces the \"black box\" risk of deep learning while providing higher accuracy and stability than traditional decision trees.",
        "Dynamic Vitals Integration: Unlike basic symptom checkers, the invention incorporates real-time physiological data (vitals and blood tests) into its predictive feature vector.",
        "Triage Efficiency: Automates the classification of patient urgency, thereby reducing the burden on primary care emergency facilities."
    ]
    for n in novelties:
        doc.add_paragraph(n, style='List Bullet')

    # --- 5. Brief Description of Drawings ---
    add_section_header("5. BRIEF DESCRIPTION OF THE DRAWINGS")
    drawings = [
        "Figure 1: System Architecture Diagram showing the interaction between the User, Frontend, Node.js Backend, and Python ML Engine.",
        "Figure 2: Operational Flowchart depicting the step-by-step logic from user login to diagnosis and subsequent care actions.",
        "Figure 3: Data Processing Pipeline showing how vitals and symptoms are converted into a feature vector for ML inference."
    ]
    for d in drawings:
        doc.add_paragraph(d, style='List Bullet')

    # --- 6. Detailed Description ---
    add_section_header("6. DETAILED DESCRIPTION OF THE INVENTION")
    add_normal_paragraph("The invention comprises several interconnected modules. The Medidose Engine is the core diagnostic component, implemented in Python using the scikit-learn library. It utilizes a pre-trained Random Forest model that has been trained on clinical data encompassing diverse symptoms and physiological markers.")
    add_normal_paragraph("When a user initiates the process via the Web Interface, the system captures data in a multi-step form. This data is transmitted via an encrypted REST API to the Backend Server (Node.js). The backend spawns a child process to invoke the ML engine, passing the data as a JSON object. The engine calculates the BMI, normalizes vital signs, and applies one-hot encoding to symptoms. The top three probable conditions are returned to the user with confidence percentages and specialist recommendations.")

    # --- 7. System Architecture ---
    add_section_header("7. SYSTEM ARCHITECTURE EXPLANATION")
    arch = [
        "Presentation Layer: HTML5/CSS3/JS frontend providing a responsive UI.",
        "Logic Layer: Node.js/Express server handling authentication (JWT), routing, and database interactions.",
        "Inference Layer: Python-based ML module performing real-time diagnosis.",
        "Data Layer: MongoDB Atlas hosting user records, doctor profiles, and appointment schedules.",
        "Integration Layer: Third-party APIs for mapping (Google Maps) and payments (Stripe)."
    ]
    for i, item in enumerate(arch, 1):
        doc.add_paragraph(f"{i}. {item}")

    # --- 8. Flowchart ---
    add_section_header("8. FLOWCHART OF OPERATION")
    steps = [
        "Start: User logs into the platform.",
        "Input: User enters vitals (BP, Heart Rate, etc.) and selects symptoms.",
        "Preprocessing: System calculates BMI and encodes categorical data.",
        "Inference: Random Forest model analyzes the feature vector.",
        "Result Display: Top 3 diseases with confidence scores and advice.",
        "Action Selection: User chooses to 'Book Appointment' or 'Analyze Insurance'.",
        "Closure: Appointment confirmed or insurance coverage verified."
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    # --- 9. Claims ---
    add_section_header("9. CLAIMS")
    claims = [
        "A method for integrated healthcare management, comprising: (a) Receiving patient vitals and symptoms via a web interface; (b) Generating a probabilistic diagnosis using a Random Forest algorithm; (c) Providing an automated specialist recommendation based on said diagnosis; (d) Facilitating an appointment booking with a recommended specialist.",
        "The method of claim 1, further comprising analyzing an insurance policy document to determine financial coverage for the generated diagnosis.",
        "The method of claim 1, wherein the vital signs include Blood Pressure, Body Temperature, and Heart Rate, which are used as numerical features in the predictive model."
    ]
    for i, claim in enumerate(claims, 1):
        doc.add_paragraph(f"{i}. {claim}")

    # --- 10. Abstract ---
    add_section_header("10. ABSTRACT")
    add_normal_paragraph("The invention provides an AI-powered healthcare accessibility platform that combines preliminary diagnosis with logistical healthcare management. By leveraging ensemble machine learning, specifically a Random Forest Classifier, the system predicts potential health conditions from a combination of user symptoms and vital signs. The platform further integrates a telemedicine booking system and insurance policy analysis, creating a unified digital pathway from initial symptom assessment to professional medical treatment and financial planning.")

    # --- 11-13. Figures ---
    add_section_header("11. FIGURE 1: SYSTEM COMPONENTS")
    doc.add_paragraph("[User] <--> [Web Browser (UI)] <--> [Express Server] <--> [ML Predictor]\n\t\t\t\t\t|\n\t\t\t\t\t+--> [MongoDB]\n\t\t\t\t\t+--> [External APIs]")

    add_section_header("12. FIGURE 2: DIAGNOSTIC LOGIC")
    doc.add_paragraph("[Vitals Input] --> [Symptom Selection] --> [Feature Encoding] --> [RF Model] --> [Top 3 Predictions]")

    add_section_header("13. FIGURE 3: POST-DIAGNOSIS ACTIONS")
    doc.add_paragraph("[Diagnosis Result]\n\t|\n\t+--> [Book Specialist] --> [Video Consultation]\n\t|\n\t+--> [Analyze Insurance] --> [Coverage Report]\n\t|\n\t+--> [Find Hospital] --> [Google Maps Routing]")

    # --- Save ---
    doc.save("PROJECT_PATENT_DRAFT.docx")
    print("Document saved successfully as PROJECT_PATENT_DRAFT.docx")

if __name__ == "__main__":
    create_patent_doc()
